"""
Word Document Builder — converts extracted PageData into a .docx file.
Preserves: fonts, sizes, bold/italic/color, alignment, tables, images.
Handles OCR for scanned pages.
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .analyzer import PageData, TextBlock, TableBlock, ImageBlock

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


ALIGN_MAP = {
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# Vertical gap (points) below which two spans are treated as the same paragraph
SAME_LINE_THRESHOLD = 4.0


class WordBuilder:
    """Builds a .docx document from a list of PageData objects."""

    def __init__(self, pages: list, ocr_lang: str = "eng"):
        self.pages = pages
        self.ocr_lang = ocr_lang
        self.doc = Document()
        self._setup_default_styles()

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def build(self, output_path: str):
        """Build and save the Word document to output_path."""
        self._clear_default_paragraph()

        for page in self.pages:
            if page.is_scanned:
                self._handle_scanned_page(page)
            else:
                self._handle_text_page(page)

            # Page break between pages (except last) — attached to last
            # paragraph so no extra blank paragraph is inserted
            if page.page_num < len(self.pages) - 1:
                self._add_page_break()

        self.doc.save(output_path)

    # ------------------------------------------------------------------
    # Page handlers
    # ------------------------------------------------------------------

    def _handle_text_page(self, page: PageData):
        """Process a text-based page: merge same-line spans, then render."""
        # Merge text blocks that share the same line into one paragraph
        merged_blocks = self._merge_same_line_blocks(page.text_blocks)

        # Collect all elements sorted top-to-bottom
        elements = []
        for tb in merged_blocks:
            elements.append(("text", tb.y0, tb))
        for table in page.tables:
            elements.append(("table", table.y0, table))
        for img in page.images:
            elements.append(("image", img.y0, img))

        elements.sort(key=lambda e: e[1])

        for kind, _, obj in elements:
            if kind == "text":
                self._add_text_block(obj)
            elif kind == "table":
                self._add_table(obj)
            elif kind == "image":
                self._add_image(obj)

    def _handle_scanned_page(self, page: PageData):
        """OCR a scanned page and add extracted text, or embed as image."""
        if not page.images:
            return

        img_block = page.images[0]

        if OCR_AVAILABLE:
            try:
                pil_img = Image.open(io.BytesIO(img_block.image_bytes))
                ocr_text = pytesseract.image_to_string(pil_img, lang=self.ocr_lang)
                if ocr_text.strip():
                    para = self.doc.add_paragraph(ocr_text.strip())
                    para.style = self.doc.styles["Normal"]
                    self._zero_spacing(para)
                    return
            except Exception:
                pass

        self._add_image(img_block)

    # ------------------------------------------------------------------
    # Same-line merging
    # ------------------------------------------------------------------

    def _merge_same_line_blocks(self, blocks: list) -> list:
        """
        Group spans whose top edges are within SAME_LINE_THRESHOLD of each
        other into a single TextBlock (joined by a space). This prevents
        every word on a line becoming its own paragraph.
        """
        if not blocks:
            return blocks

        # Sort by top-y first, then left-x
        sorted_blocks = sorted(blocks, key=lambda b: (round(b.y0), b.x0))
        groups = []
        current_group = [sorted_blocks[0]]

        for block in sorted_blocks[1:]:
            prev = current_group[-1]
            # Same line if y0 values are close
            if abs(block.y0 - prev.y0) <= SAME_LINE_THRESHOLD:
                current_group.append(block)
            else:
                groups.append(current_group)
                current_group = [block]
        groups.append(current_group)

        merged = []
        for group in groups:
            if len(group) == 1:
                merged.append(group[0])
                continue
            # Use the properties of the first (leftmost) span for the merged block
            lead = group[0]
            combined_text = " ".join(b.text for b in group if b.text)
            new_block = TextBlock(
                text=combined_text,
                x0=lead.x0, y0=lead.y0,
                x1=group[-1].x1, y1=max(b.y1 for b in group),
                font_name=lead.font_name,
                font_size=lead.font_size,
                bold=lead.bold,
                italic=lead.italic,
                color=lead.color,
                align=lead.align,
                page_num=lead.page_num,
            )
            merged.append(new_block)

        return merged

    # ------------------------------------------------------------------
    # Text block rendering
    # ------------------------------------------------------------------

    def _add_text_block(self, tb: TextBlock):
        """Add a styled paragraph for a text span — no extra spacing."""
        heading_level = self._detect_heading(tb)

        if heading_level:
            para = self.doc.add_heading("", level=heading_level)
        else:
            para = self.doc.add_paragraph()
            para.style = self.doc.styles["Normal"]

        # ── Kill all spacing around the paragraph ──────────────────────
        self._zero_spacing(para)

        para.alignment = ALIGN_MAP.get(tb.align, WD_ALIGN_PARAGRAPH.LEFT)

        run = para.add_run(tb.text)
        run.bold = tb.bold
        run.italic = tb.italic
        run.font.name = self._safe_font(tb.font_name)
        run.font.size = Pt(tb.font_size)

        r, g, b = tb.color
        if (r, g, b) != (0, 0, 0):
            run.font.color.rgb = RGBColor(r, g, b)

    def _detect_heading(self, tb: TextBlock):
        """Return heading level (int) if text looks like a heading, else None."""
        if tb.font_size >= 24 and tb.bold:
            return 1
        if tb.font_size >= 18 and tb.bold:
            return 2
        if tb.font_size >= 14 and tb.bold:
            return 3
        if tb.bold and tb.font_size >= 13:
            return 4
        return None

    # ------------------------------------------------------------------
    # Table rendering
    # ------------------------------------------------------------------

    def _add_table(self, table_block: TableBlock):
        """Render a TableBlock as a Word table — no trailing blank paragraph."""
        rows = table_block.rows
        if not rows:
            return

        num_cols = max(len(row) for row in rows)
        if num_cols == 0:
            return

        word_table = self.doc.add_table(rows=len(rows), cols=num_cols)
        word_table.style = "Table Grid"

        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                if c_idx >= num_cols:
                    break
                cell = word_table.cell(r_idx, c_idx)
                cell.text = cell_text or ""
                for para in cell.paragraphs:
                    self._zero_spacing(para)
                    for run in para.runs:
                        run.font.size = Pt(10)

        # No self.doc.add_paragraph() here — that was adding the blank gap

    # ------------------------------------------------------------------
    # Image rendering
    # ------------------------------------------------------------------

    def _add_image(self, img_block: ImageBlock):
        """Embed an image — no extra spacing paragraph."""
        try:
            img_stream = io.BytesIO(img_block.image_bytes)

            width_in  = img_block.width_pt  / 72.0
            height_in = img_block.height_pt / 72.0

            max_width = 6.5
            if width_in > max_width:
                scale     = max_width / width_in
                width_in  = max_width
                height_in *= scale

            para = self.doc.add_paragraph()
            self._zero_spacing(para)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(img_stream, width=Inches(width_in))
        except Exception:
            pass

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _zero_spacing(para):
        """Remove space_before and space_after from a paragraph."""
        fmt = para.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after  = Pt(0)

    def _setup_default_styles(self):
        """Set global document defaults — zero spacing, standard font."""
        style = self.doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        # Zero out paragraph spacing on Normal style globally
        fmt = style.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after  = Pt(0)

        section = self.doc.sections[0]
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    def _clear_default_paragraph(self):
        """Remove the blank paragraph python-docx inserts by default."""
        for para in self.doc.paragraphs:
            if para.text == "":
                p = para._element
                p.getparent().remove(p)

    def _add_page_break(self):
        """Attach a page break to the last paragraph — no extra blank line."""
        paras = self.doc.paragraphs
        if paras:
            # Re-use last paragraph's last run to carry the page break
            last_para = paras[-1]
            run = last_para.add_run()
        else:
            last_para = self.doc.add_paragraph()
            run = last_para.add_run()

        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)

    @staticmethod
    def _safe_font(font_name: str) -> str:
        """Map PDF font names to Word-safe fonts."""
        n = font_name.lower()
        if "times"    in n or "serif"     in n: return "Times New Roman"
        if "courier"  in n or "mono"      in n: return "Courier New"
        if "arial"    in n or "helvetica" in n or "sans" in n: return "Arial"
        if "calibri"  in n: return "Calibri"
        if "georgia"  in n: return "Georgia"
        return "Arial"


# Re-export so analyzer imports still work
from .analyzer import TextBlock  # noqa: E402, F401
